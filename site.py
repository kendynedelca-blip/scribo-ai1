import base64
import hashlib
import hmac
import sqlite3
import json
import asyncio
import re
import secrets
from datetime import datetime
from io import BytesIO

import streamlit as st
import streamlit.components.v1 as components
from groq import Groq

# IMPORTANT: Streamlit cere ca set_page_config să fie prima comandă Streamlit.
st.set_page_config(
    page_title="Scribo AI",
    page_icon="S",
    layout="wide",
    initial_sidebar_state="auto",
)

# Ascundem doar elementele decorative, nu întregul header.
# Headerul păstrează controlul sidebarului pe mobil și accesibilitatea tastaturii.
st.markdown(
    """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    [data-testid="stStatusWidget"] {visibility: hidden;}
    .viewerBadge_container__1QS13 {display: none !important;}
    header[data-testid="stHeader"] {
        background: transparent !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ============================================================
# DEPENDENȚE OPȚIONALE
# ============================================================

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import edge_tts
    EDGE_TTS_AVAILABLE = True
except ImportError:
    EDGE_TTS_AVAILABLE = False

try:
    from tavily import TavilyClient
    TAVILY_AVAILABLE = True
except ImportError:
    TAVILY_AVAILABLE = False



MODEL_GROQ = "openai/gpt-oss-120b"
WHISPER_MODEL = "whisper-large-v3-turbo"

DB_FILE = "users_data.db"


# ============================================================
# UTILITARE
# ============================================================

def enhance_accessibility_and_follow_chat(scroll_to_bottom=False):
    """Etichetează controalele importante și urmărește ultimul mesaj fără a fura focusul."""
    scroll_js = "true" if scroll_to_bottom else "false"
    components.html(
        f"""
        <script>
        (() => {{
          const doc = window.parent.document;
          const label = (selector, text) => {{
            const el = doc.querySelector(selector);
            if (el && !el.getAttribute('aria-label')) el.setAttribute('aria-label', text);
          }};
          label('[data-testid="stSidebarCollapsedControl"] button', 'Deschide meniul lateral');
          label('[data-testid="collapsedControl"] button', 'Deschide meniul lateral');
          label('[data-testid="stChatInput"] textarea', 'Mesaj pentru Scribo');

          if ({scroll_js}) {{
            const anchor = doc.getElementById('chat-bottom-anchor');
            if (anchor) {{
              anchor.scrollIntoView({{behavior: 'auto', block: 'end'}});
            }}
          }}
        }})();
        </script>
        """,
        height=0,
        width=0,
    )


def now_iso():
    return datetime.now().isoformat(timespec="seconds")


def make_id():
    return secrets.token_hex(16)


def escape_html(text):
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def clean_ai_text(text):
    cleaned = str(text or "")
    cleaned = re.sub(r"```(?:\w+)?\s*", "", cleaned)
    cleaned = cleaned.replace("```", "")
    cleaned = re.sub(r"^\s{0,3}#{1,6}\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\*{1,3}", "", cleaned)
    cleaned = re.sub(r"_{1,3}", "", cleaned)
    cleaned = re.sub(r"^\s*[-+>]\s+", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^\s*\d+[.)]\s+", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", cleaned)
    return cleaned.strip()


def render_suggestions():
    suggestions = [
        ("✍️ Scenă", "Scrie următoarea scenă literară în stilul selectat."),
        ("🧠 Personaj", "Creează profilul psihologic al unui nou personaj."),
        ("💡 Intrigă", "Propune o răsturnare de situație neașteptată."),
        ("☕ Decor", "Descrie detaliat atmosfera și cadrul scenei."),
        ("✨ Ghidare", "Analizează stadiul curent al romanului și propune direcții de continuare."),
    ]

    suggestions_row = st.container()
    suggestion_cols = suggestions_row.columns(5)

    for column, (label, pending_prompt) in zip(suggestion_cols, suggestions):
        with column:
            if st.button(label, use_container_width=True):
                st.session_state.pending_prompt = pending_prompt
                st.rerun()


# ============================================================
# PAROLE
# ============================================================

def hash_pw(password: str) -> str:
    """
    Parolele noi sunt stocate cu PBKDF2 + salt.
    """
    salt = secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt,
        200_000,
    )

    return (
        "pbkdf2_sha256$200000$"
        + base64.b64encode(salt).decode("ascii")
        + "$"
        + base64.b64encode(digest).decode("ascii")
    )


def verify_pw(password: str, stored_hash: str) -> bool:
    if not stored_hash:
        return False

    # Format nou
    if stored_hash.startswith("pbkdf2_sha256$"):
        try:
            _, iterations, salt_b64, digest_b64 = stored_hash.split("$")

            salt = base64.b64decode(salt_b64)
            expected = base64.b64decode(digest_b64)

            actual = hashlib.pbkdf2_hmac(
                "sha256",
                password.encode("utf-8"),
                salt,
                int(iterations),
            )

            return hmac.compare_digest(actual, expected)
        except Exception:
            return False

    # Compatibilitate cu parolele vechi SHA256
    old_hash = hashlib.sha256(
        password.encode("utf-8")
    ).hexdigest()

    return hmac.compare_digest(old_hash, stored_hash)


# ============================================================
# BAZĂ DE DATE
# ============================================================

def get_conn():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn


def column_exists(conn, table_name, column_name):
    rows = conn.execute(
        f"PRAGMA table_info({table_name})"
    ).fetchall()

    return any(row["name"] == column_name for row in rows)


def add_column_if_missing(conn, table_name, column_name, definition):
    if not column_exists(conn, table_name, column_name):
        conn.execute(
            f"ALTER TABLE {table_name} ADD COLUMN {column_name} {definition}"
        )


def init_db():
    conn = get_conn()
    c = conn.cursor()

    # --------------------------------------------------------
    # TABEL USERS - păstrăm compatibilitatea cu baza veche
    # --------------------------------------------------------

    c.execute(
        """
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT NOT NULL,
            messages TEXT,
            roman TEXT,
            style TEXT,
            voice TEXT
        )
        """
    )

    # Coloane noi
    add_column_if_missing(conn, "users", "email", "TEXT")
    add_column_if_missing(conn, "users", "display_name", "TEXT")
    add_column_if_missing(conn, "users", "auth_provider", "TEXT")
    add_column_if_missing(conn, "users", "provider_sub", "TEXT")
    add_column_if_missing(conn, "users", "created_at", "TEXT")

    # --------------------------------------------------------
    # CONVERSAȚII
    # --------------------------------------------------------

    c.execute(
        """
        CREATE TABLE IF NOT EXISTS conversations (
            id TEXT PRIMARY KEY,
            username TEXT NOT NULL,
            title TEXT NOT NULL,
            project_id TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )

    # --------------------------------------------------------
    # MESAJE
    # --------------------------------------------------------

    c.execute(
        """
        CREATE TABLE IF NOT EXISTS chat_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id TEXT NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            images TEXT,
            created_at TEXT NOT NULL
        )
        """
    )

    # --------------------------------------------------------
    # PROIECTE
    # --------------------------------------------------------

    c.execute(
        """
        CREATE TABLE IF NOT EXISTS projects (
            id TEXT PRIMARY KEY,
            username TEXT NOT NULL,
            name TEXT NOT NULL,
            roman TEXT DEFAULT '',
            style TEXT DEFAULT 'Fantezie Alchimie',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )

    # --------------------------------------------------------
    # INDEXURI
    # --------------------------------------------------------

    c.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_conversations_user
        ON conversations(username)
        """
    )

    c.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_messages_conversation
        ON chat_messages(conversation_id)
        """
    )

    c.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_projects_user
        ON projects(username)
        """
    )

    conn.commit()

    # --------------------------------------------------------
    # MIGRAREA DATELOR VECHI
    # --------------------------------------------------------

    migrate_old_data(conn)

    conn.commit()
    conn.close()


def migrate_old_data(conn):
    """
    Transformă datele vechi:
      users.messages -> conversations + chat_messages
      users.roman    -> projects
    """

    old_users = conn.execute(
        """
        SELECT username, password, messages, roman, style, voice,
               email, display_name, auth_provider, provider_sub, created_at
        FROM users
        """
    ).fetchall()

    for user in old_users:
        username = user["username"]

        if not user["created_at"]:
            conn.execute(
                """
                UPDATE users
                SET created_at = ?
                WHERE username = ?
                """,
                (now_iso(), username),
            )

        if not user["display_name"]:
            conn.execute(
                """
                UPDATE users
                SET display_name = ?
                WHERE username = ?
                """,
                (username, username),
            )

        # ----------------------------------------------------
        # PROIECT VECHI
        # ----------------------------------------------------

        existing_project = conn.execute(
            """
            SELECT id
            FROM projects
            WHERE username = ?
            ORDER BY created_at ASC
            LIMIT 1
            """,
            (username,),
        ).fetchone()

        if existing_project:
            project_id = existing_project["id"]
        else:
            project_id = make_id()

            roman = user["roman"] or ""
            style = user["style"] or "Fantezie Alchimie"

            conn.execute(
                """
                INSERT INTO projects
                (id, username, name, roman, style, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    project_id,
                    username,
                    "Romanul meu",
                    roman,
                    style,
                    now_iso(),
                    now_iso(),
                ),
            )

        # ----------------------------------------------------
        # CONVERSAȚIA VECHE
        # ----------------------------------------------------

        existing_conversation = conn.execute(
            """
            SELECT id
            FROM conversations
            WHERE username = ?
            LIMIT 1
            """,
            (username,),
        ).fetchone()

        if existing_conversation:
            continue

        raw_messages = user["messages"] or "[]"

        try:
            old_messages = json.loads(raw_messages)
        except Exception:
            old_messages = []

        if not old_messages:
            continue

        conversation_id = make_id()

        conn.execute(
            """
            INSERT INTO conversations
            (id, username, title, project_id, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                conversation_id,
                username,
                "Conversația mea",
                project_id,
                now_iso(),
                now_iso(),
            ),
        )

        for msg in old_messages:
            role = msg.get("role", "")
            content = msg.get("content", "")

            if role not in ("user", "assistant"):
                continue

            images = msg.get("images", [])

            # Audio nu îl punem în DB
            try:
                clean_images = []

                for image in images:
                    clean_images.append(
                        {
                            "name": image.get("name", "imagine"),
                            "data": image.get("data", b""),
                        }
                    )

                images_json = json.dumps(
                    clean_images,
                    default=str,
                )
            except Exception:
                images_json = "[]"

            conn.execute(
                """
                INSERT INTO chat_messages
                (conversation_id, role, content, images, created_at)
                VALUES (?, ?, ?, ?, ?)
                """,
                (
                    conversation_id,
                    role,
                    content,
                    images_json,
                    now_iso(),
                ),
            )


init_db()


# ============================================================
# USERS
# ============================================================

def register_user(username, email, password):
    username = username.strip()
    email = email.strip().lower()

    if not username or not email or not password:
        return False, "Completează toate câmpurile."

    if len(username) < 3:
        return False, "Numele de utilizator trebuie să aibă cel puțin 3 caractere."

    if len(password) < 6:
        return False, "Parola trebuie să aibă cel puțin 6 caractere."

    conn = get_conn()

    try:
        existing_username = conn.execute(
            """
            SELECT username
            FROM users
            WHERE LOWER(username) = LOWER(?)
            """,
            (username,),
        ).fetchone()

        if existing_username:
            return False, "Numele de utilizator este deja folosit."

        existing_email = conn.execute(
            """
            SELECT username
            FROM users
            WHERE LOWER(email) = LOWER(?)
            """,
            (email,),
        ).fetchone()

        if existing_email:
            return False, "Emailul este deja folosit."

        conn.execute(
            """
            INSERT INTO users
            (
                username,
                password,
                messages,
                roman,
                style,
                voice,
                email,
                display_name,
                auth_provider,
                provider_sub,
                created_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                username,
                hash_pw(password),
                "[]",
                "",
                "Fantezie Alchimie",
                "ro-RO-AlinaNeural",
                email,
                username,
                "local",
                "",
                now_iso(),
            ),
        )

        conn.commit()

        # proiect implicit
        create_project(
            username,
            "Romanul meu",
            "Fantezie Alchimie",
        )

        return True, "Cont creat cu succes."

    except sqlite3.IntegrityError:
        return False, "Nu am putut crea contul."

    finally:
        conn.close()


def login_user(username_or_email, password):
    value = username_or_email.strip()

    conn = get_conn()

    row = conn.execute(
        """
        SELECT *
        FROM users
        WHERE LOWER(username) = LOWER(?)
           OR LOWER(email) = LOWER(?)
        LIMIT 1
        """,
        (value, value),
    ).fetchone()

    conn.close()

    if not row:
        return None

    if not verify_pw(password, row["password"]):
        return None

    # Dacă era parolă veche SHA256, o actualizăm automat.
    if not row["password"].startswith("pbkdf2_sha256$"):
        conn = get_conn()

        conn.execute(
            """
            UPDATE users
            SET password = ?
            WHERE username = ?
            """,
            (
                hash_pw(password),
                row["username"],
            ),
        )

        conn.commit()
        conn.close()

    return {
        "username": row["username"],
        "email": row["email"] or "",
        "display_name": row["display_name"] or row["username"],
        "style": row["style"] or "Fantezie Alchimie",
        "voice": row["voice"] or "ro-RO-AlinaNeural",
    }


def get_user(username):
    conn = get_conn()

    row = conn.execute(
        """
        SELECT *
        FROM users
        WHERE username = ?
        """,
        (username,),
    ).fetchone()

    conn.close()

    return row


def get_user_by_provider(provider, provider_sub):
    if not provider_sub:
        return None

    conn = get_conn()

    row = conn.execute(
        """
        SELECT *
        FROM users
        WHERE auth_provider = ?
          AND provider_sub = ?
        LIMIT 1
        """,
        (provider, provider_sub),
    ).fetchone()

    conn.close()

    return row


def create_or_get_oauth_user(provider, provider_data):
    provider_sub = str(
        provider_data.get("sub", "")
    ).strip()

    email = str(
        provider_data.get("email", "")
    ).strip().lower()

    name = str(
        provider_data.get("name", "")
        or provider_data.get("given_name", "")
        or "Utilizator"
    ).strip()

    if not provider_sub:
        return None

    existing = get_user_by_provider(
        provider,
        provider_sub,
    )

    if existing:
        return existing["username"]

    conn = get_conn()

    # Dacă emailul există deja, îl legăm numai când providerul confirmă emailul.
    # Unele OIDC-uri nu trimit email_verified; în acest caz nu facem auto-link.
    email_verified = provider_data.get("email_verified")
    verified = email_verified is True or str(email_verified).lower() == "true"

    if email and verified:
        existing_email = conn.execute(
            """
            SELECT username
            FROM users
            WHERE LOWER(email) = LOWER(?)
            LIMIT 1
            """,
            (email,),
        ).fetchone()

        if existing_email:
            username = existing_email["username"]

            conn.execute(
                """
                UPDATE users
                SET
                    auth_provider = ?,
                    provider_sub = ?,
                    display_name = ?
                WHERE username = ?
                """,
                (
                    provider,
                    provider_sub,
                    name,
                    username,
                ),
            )

            conn.commit()
            conn.close()

            return username

    # Username generat
    base_username = "".join(
        ch.lower()
        for ch in name
        if ch.isalnum()
    )[:20]

    if not base_username:
        base_username = "utilizator"

    username = base_username
    counter = 1

    while True:
        exists = conn.execute(
            """
            SELECT username
            FROM users
            WHERE LOWER(username) = LOWER(?)
            """,
            (username,),
        ).fetchone()

        if not exists:
            break

        counter += 1
        username = f"{base_username}{counter}"

    conn.execute(
        """
        INSERT INTO users
        (
            username,
            password,
            messages,
            roman,
            style,
            voice,
            email,
            display_name,
            auth_provider,
            provider_sub,
            created_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            username,
            hash_pw(secrets.token_urlsafe(32)),
            "[]",
            "",
            "Fantezie Alchimie",
            "ro-RO-AlinaNeural",
            email,
            name,
            provider,
            provider_sub,
            now_iso(),
        ),
    )

    conn.commit()
    conn.close()

    create_project(
        username,
        "Romanul meu",
        "Fantezie Alchimie",
    )

    return username


# ============================================================
# PROIECTE
# ============================================================

def create_project(username, name, style="Fantezie Alchimie"):
    project_id = make_id()

    conn = get_conn()

    conn.execute(
        """
        INSERT INTO projects
        (id, username, name, roman, style, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            project_id,
            username,
            name.strip() or "Proiect nou",
            "",
            style,
            now_iso(),
            now_iso(),
        ),
    )

    conn.commit()
    conn.close()

    return project_id


def get_projects(username):
    conn = get_conn()

    rows = conn.execute(
        """
        SELECT *
        FROM projects
        WHERE username = ?
        ORDER BY updated_at DESC
        """,
        (username,),
    ).fetchall()

    conn.close()

    return rows


def get_project(project_id, username):
    conn = get_conn()

    row = conn.execute(
        """
        SELECT *
        FROM projects
        WHERE id = ?
          AND username = ?
        """,
        (
            project_id,
            username,
        ),
    ).fetchone()

    conn.close()

    return row


def update_project(project_id, username, roman, style):
    conn = get_conn()

    conn.execute(
        """
        UPDATE projects
        SET roman = ?,
            style = ?,
            updated_at = ?
        WHERE id = ?
          AND username = ?
        """,
        (
            roman,
            style,
            now_iso(),
            project_id,
            username,
        ),
    )

    conn.commit()
    conn.close()


# ============================================================
# CONVERSAȚII
# ============================================================

def create_conversation(username, project_id=None, title="Conversație nouă"):
    conversation_id = make_id()

    conn = get_conn()

    conn.execute(
        """
        INSERT INTO conversations
        (id, username, title, project_id, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            conversation_id,
            username,
            title,
            project_id,
            now_iso(),
            now_iso(),
        ),
    )

    conn.commit()
    conn.close()

    return conversation_id


def get_conversations(username):
    conn = get_conn()

    rows = conn.execute(
        """
        SELECT *
        FROM conversations
        WHERE username = ?
        ORDER BY updated_at DESC
        """,
        (username,),
    ).fetchall()

    conn.close()

    return rows


def get_conversation(conversation_id, username):
    conn = get_conn()

    row = conn.execute(
        """
        SELECT *
        FROM conversations
        WHERE id = ?
          AND username = ?
        """,
        (
            conversation_id,
            username,
        ),
    ).fetchone()

    conn.close()

    return row


def get_conversation_messages(conversation_id, username):
    conversation = get_conversation(
        conversation_id,
        username,
    )

    if not conversation:
        return []

    conn = get_conn()

    rows = conn.execute(
        """
        SELECT *
        FROM chat_messages
        WHERE conversation_id = ?
        ORDER BY id ASC
        """,
        (conversation_id,),
    ).fetchall()

    conn.close()

    messages = []

    for row in rows:
        try:
            images = json.loads(row["images"] or "[]")
        except Exception:
            images = []

        messages.append(
            {
                "role": row["role"],
                "content": row["content"],
                "images": images,
            }
        )

    return messages


def add_message(
    conversation_id,
    username,
    role,
    content,
    images=None,
):
    conversation = get_conversation(
        conversation_id,
        username,
    )

    if not conversation:
        return

    images = images or []

    clean_images = []

    for image in images:
        try:
            clean_images.append(
                {
                    "name": image.get("name", "imagine"),
                    "data": image.get("data", b""),
                }
            )
        except Exception:
            pass

    try:
        images_json = json.dumps(
            clean_images,
            default=str,
        )
    except Exception:
        images_json = "[]"

    conn = get_conn()

    conn.execute(
        """
        INSERT INTO chat_messages
        (conversation_id, role, content, images, created_at)
        VALUES (?, ?, ?, ?, ?)
        """,
        (
            conversation_id,
            role,
            content,
            images_json,
            now_iso(),
        ),
    )

    conn.execute(
        """
        UPDATE conversations
        SET updated_at = ?
        WHERE id = ?
          AND username = ?
        """,
        (
            now_iso(),
            conversation_id,
            username,
        ),
    )

    conn.commit()
    conn.close()


def update_conversation_title(
    conversation_id,
    username,
    title,
):
    conn = get_conn()

    conn.execute(
        """
        UPDATE conversations
        SET title = ?,
            updated_at = ?
        WHERE id = ?
          AND username = ?
        """,
        (
            title[:80],
            now_iso(),
            conversation_id,
            username,
        ),
    )

    conn.commit()
    conn.close()


def delete_conversation(conversation_id, username):
    conn = get_conn()

    conn.execute(
        """
        DELETE FROM chat_messages
        WHERE conversation_id = ?
        """,
        (conversation_id,),
    )

    conn.execute(
        """
        DELETE FROM conversations
        WHERE id = ?
          AND username = ?
        """,
        (
            conversation_id,
            username,
        ),
    )

    conn.commit()
    conn.close()


def make_conversation_title(prompt):
    clean = " ".join(prompt.strip().split())

    if not clean:
        return "Conversație nouă"

    # Dacă este foarte lung, îl scurtăm
    if len(clean) > 55:
        clean = clean[:55].rsplit(" ", 1)[0] + "..."

    return clean[0].upper() + clean[1:]


# ============================================================
# CLIENT GROQ
# ============================================================

try:
    API_KEY = st.secrets["CHEIE_GROQ"]
except Exception:
    st.error(
        "Cheia Groq nu a fost găsită în "
        ".streamlit/secrets.toml sub numele CHEIE_GROQ."
    )
    st.stop()

client = Groq(api_key=API_KEY)

TAVILY_KEY = st.secrets.get(
    "TAVILY_API_KEY",
    "",
)

tavily_client = (
    TavilyClient(api_key=TAVILY_KEY)
    if (TAVILY_AVAILABLE and TAVILY_KEY)
    else None
)


# ============================================================
# SESSION STATE
# ============================================================

defaults = {
    "authenticated": False,
    "user": None,
    "username": None,
    "email": "",
    "display_name": "",
    "roman": "",
    "messages": [],
    "style": "Fantezie Alchimie",
    "voice_model": "ro-RO-AlinaNeural",
    "attached_text": "",
    "attached_file_name": "",
    "attached_images": [],
    "pending_prompt": None,
    "is_voice_input": False,
    "conversation_id": None,
    "project_id": None,
    "auth_initialized": False,
    "show_login": True,
}

for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value


# ============================================================
# OIDC - GOOGLE / APPLE
# ============================================================

def oidc_available():
    try:
        return bool(
            hasattr(st, "user")
            and hasattr(st.user, "is_logged_in")
        )
    except Exception:
        return False


def _secret_section(name):
    """Returnează în siguranță o secțiune din st.secrets ca dicționar."""
    try:
        value = st.secrets.get(name, {})
        if hasattr(value, "to_dict"):
            return value.to_dict()
        return dict(value)
    except Exception:
        return {}


def oidc_provider_configured(provider):
    """Verifică providerul numit conform formatului oficial [auth.provider]."""
    auth = _secret_section("auth")
    try:
        provider_cfg = auth.get(provider, {})
        if hasattr(provider_cfg, "to_dict"):
            provider_cfg = provider_cfg.to_dict()
        else:
            provider_cfg = dict(provider_cfg)
    except Exception:
        provider_cfg = {}

    shared_ok = bool(auth.get("redirect_uri") and auth.get("cookie_secret"))
    provider_ok = all(
        provider_cfg.get(key)
        for key in ("client_id", "client_secret", "server_metadata_url")
    )
    return shared_ok and provider_ok


def oidc_default_configured():
    auth = _secret_section("auth")
    return all(
        auth.get(key)
        for key in (
            "redirect_uri",
            "cookie_secret",
            "client_id",
            "client_secret",
            "server_metadata_url",
        )
    )


def start_oidc_login(provider=None):
    """Pornește OIDC fără a lăsa aplicația într-o stare intermediară."""
    try:
        if provider:
            st.login(provider)
        else:
            st.login()
    except Exception as exc:
        st.error(
            "Conectarea externă nu a putut fi pornită. "
            "Verifică redirect_uri, datele providerului și pachetul Authlib."
        )
        st.caption(f"Detaliu tehnic: {type(exc).__name__}")


def handle_oidc_login():
    """
    Dacă utilizatorul s-a autentificat prin Google/Apple,
    îl legăm de contul intern SQLite.
    """

    if not oidc_available():
        return

    try:
        if not st.user.is_logged_in:
            return
    except Exception:
        return

    try:
        user_data = st.user.to_dict()
    except Exception:
        user_data = dict(st.user)

    issuer = str(user_data.get("iss", "")).lower()
    provider_hint = str(user_data.get("provider", "")).lower()

    if "google" in issuer or provider_hint == "google":
        provider_name = "google"
    elif "apple" in issuer or "appleid.apple.com" in issuer or provider_hint == "apple":
        provider_name = "apple"
    else:
        provider_name = "oidc"

    username = create_or_get_oauth_user(
        provider_name,
        user_data,
    )

    if not username:
        st.error(
            "Autentificarea externă a reușit, "
            "dar nu am putut crea contul intern."
        )
        st.stop()

    user = get_user(username)

    if not user:
        st.error("Contul nu a putut fi încărcat.")
        st.stop()

    st.session_state.authenticated = True
    st.session_state.user = username
    st.session_state.username = username
    st.session_state.email = user["email"] or ""
    st.session_state.display_name = (
        user["display_name"] or username
    )
    st.session_state.style = (
        user["style"] or "Fantezie Alchimie"
    )
    st.session_state.voice_model = (
        user["voice"] or "ro-RO-AlinaNeural"
    )

    if not st.session_state.auth_initialized:
        start_new_conversation()
        st.session_state.auth_initialized = True


# ============================================================
# FUNCȚII DE ÎNCĂRCARE
# ============================================================

def load_conversation(conversation_id):
    username = st.session_state.username

    conversation = get_conversation(
        conversation_id,
        username,
    )

    if not conversation:
        return False

    st.session_state.conversation_id = conversation_id
    st.session_state.project_id = conversation["project_id"]

    messages = get_conversation_messages(
        conversation_id,
        username,
    )

    st.session_state.messages = messages

    if conversation["project_id"]:
        project = get_project(
            conversation["project_id"],
            username,
        )

        if project:
            st.session_state.roman = project["roman"] or ""
            st.session_state.style = (
                project["style"]
                or "Fantezie Alchimie"
            )

    return True


def start_new_conversation():
    username = st.session_state.username

    projects = get_projects(username)

    if projects:
        project_id = projects[0]["id"]
        st.session_state.project_id = project_id
        st.session_state.roman = projects[0]["roman"] or ""
        st.session_state.style = (
            projects[0]["style"]
            or "Fantezie Alchimie"
        )
    else:
        project_id = create_project(
            username,
            "Romanul meu",
            st.session_state.style,
        )

        st.session_state.project_id = project_id
        st.session_state.roman = ""

    conversation_id = create_conversation(
        username,
        project_id,
        "Conversație nouă",
    )

    st.session_state.conversation_id = conversation_id
    st.session_state.messages = []
    st.session_state.attached_text = ""
    st.session_state.attached_file_name = ""
    st.session_state.attached_images = []
    st.session_state.is_voice_input = False


# ============================================================
# CĂUTARE WEB
# ============================================================

def needs_web_search(prompt_text):
    keywords = [
        "carte",
        "autor",
        "publicat",
        "apărut",
        "2023",
        "2024",
        "2025",
        "2026",
        "intriga",
        "sinopsis",
        "editura",
        "recenzie",
        "realitate",
        "istoric",
        "cine a scris",
        "cine este",
        "despre ce e vorba in",
        "despre ce e vorba în",
        "povestea din cartea",
    ]

    prompt_lower = prompt_text.lower()

    return any(
        keyword in prompt_lower
        for keyword in keywords
    )


def get_web_context(query):
    if not tavily_client:
        return ""

    try:
        search_result = tavily_client.search(
            query=query,
            search_depth="advanced",
            max_results=3,
        )

        context_snippets = [
            f"Sursă ({res.get('url', '')}):\n"
            f"{res.get('content', '')}"
            for res in search_result.get(
                "results",
                [],
            )
        ]

        return "\n\n".join(
            context_snippets
        )

    except Exception:
        return ""


# ============================================================
# DOCUMENTE
# ============================================================

def extract_text_from_txt(file):
    try:
        return (
            file.read()
            .decode("utf-8", errors="ignore")
            .strip()
        )
    except Exception:
        return ""


def extract_text_from_pdf(file):
    if not PDF_AVAILABLE:
        return ""

    try:
        reader = PdfReader(file)

        parts = []

        for page in reader.pages:
            text = page.extract_text()

            if text:
                parts.append(text)

        return "\n\n".join(parts).strip()

    except Exception:
        return ""


def extract_text_from_docx(file):
    if not DOCX_AVAILABLE:
        return ""

    try:
        document = Document(file)

        parts = [
            p.text
            for p in document.paragraphs
            if p.text.strip()
        ]

        return "\n".join(parts).strip()

    except Exception:
        return ""


def extract_document_text(uploaded_file):
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return extract_text_from_pdf(
            uploaded_file
        )

    if name.endswith(".docx"):
        return extract_text_from_docx(
            uploaded_file
        )

    return extract_text_from_txt(
        uploaded_file
    )


def image_to_data_uri(
    img_bytes,
    mime="image/png",
):
    if isinstance(img_bytes, str):
        return img_bytes

    b64 = base64.b64encode(
        img_bytes
    ).decode("utf-8")

    return f"data:{mime};base64,{b64}"


# ============================================================
# VOCE
# ============================================================

async def generate_voice_bytes(text, voice):
    clean_text = clean_ai_text(text)[:500]

    communicate = edge_tts.Communicate(
        clean_text,
        voice,
    )

    audio_buffer = bytearray()

    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            audio_buffer.extend(
                chunk["data"]
            )

    return bytes(audio_buffer)


def get_audio_response(text, voice):
    if (
        not EDGE_TTS_AVAILABLE
        or not text.strip()
    ):
        return None

    try:
        return asyncio.run(
            generate_voice_bytes(
                text,
                voice,
            )
        )
    except Exception:
        return None


# ============================================================
# CSS - TEMA SCRIBO AI
# ============================================================

st.markdown(
    """
    <style>

    @import url(
        'https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap'
    );

    :root {
        --scribo-bg: #1a1917;
        --scribo-panel: #23221f;
        --scribo-input: #242220;
        --scribo-input-hover: #2b2926;
        --scribo-border: #3d3934;
        --scribo-text: #f1f1f1;
        --scribo-muted: #aaa39a;
        --scribo-accent: #d9774a;
    }

    /* ========================================================
       FUNDAL PRINCIPAL
       ======================================================== */

    html,
    body,
    .stApp,
    [data-testid="stAppViewContainer"],
    [data-testid="stMain"] {
        background-color: var(--scribo-bg) !important;
        color: var(--scribo-text) !important;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
    }

    footer {
        visibility: hidden;
    }

    header[data-testid="stHeader"] {
        visibility: visible;
        background: transparent !important;
    }

    /* ========================================================
       TEXT GENERAL
       ======================================================== */

    .stApp p,
    .stApp label,
    .stApp h1,
    .stApp h2,
    .stApp h3,
    .stApp h4,
    .stApp h5,
    .stApp h6 {
        color: #f1f1f1;
    }

    .stApp [data-testid="stMarkdownContainer"] p,
    .stApp [data-testid="stMarkdownContainer"] span {
        color: #f1f1f1;
    }

    /* ========================================================
       LABEL-URI
       ======================================================== */

    .stApp [data-testid="stTextInput"] label,
    .stApp [data-testid="stTextArea"] label,
    .stApp [data-testid="stSelectbox"] label,
    .stApp [data-testid="stFileUploader"] label,
    .stApp [data-testid="stRadio"] label {
        color: #ffffff !important;
    }

    /* ========================================================
       INPUTURI
       ======================================================== */

    .stApp input,
    .stApp textarea {
        background-color: #242220 !important;
        color: #ffffff !important;
        -webkit-text-fill-color: #ffffff !important;
        caret-color: #d9774a !important;
        border: 1px solid #3d3934 !important;
        border-radius: 10px !important;
        box-shadow: none !important;
    }

    .stApp input:hover,
    .stApp textarea:hover,
    .stApp input:focus,
    .stApp textarea:focus {
        background-color: #2b2926 !important;
        color: #ffffff !important;
        -webkit-text-fill-color: #ffffff !important;
        border-color: #5a5149 !important;
        box-shadow: 0 0 0 1px #5a5149 !important;
    }

    /* PLACEHOLDER */

    .stApp input::placeholder,
    .stApp textarea::placeholder {
        color: #aaa39a !important;
        -webkit-text-fill-color: #aaa39a !important;
        opacity: 1 !important;
    }

    /* ========================================================
       AUTOFILL - IMPORTANT PENTRU PAROLE
       ======================================================== */

    .stApp input:-webkit-autofill,
    .stApp input:-webkit-autofill:hover,
    .stApp input:-webkit-autofill:focus {
        -webkit-text-fill-color: #ffffff !important;
        -webkit-box-shadow:
            0 0 0 1000px #242220 inset !important;
        box-shadow:
            0 0 0 1000px #242220 inset !important;
        caret-color: #ffffff !important;
    }

    /* ========================================================
       BASEWEB INPUT
       ======================================================== */

    .stApp [data-baseweb="input"],
    .stApp [data-baseweb="base-input"],
    .stApp [data-baseweb="textarea"],
    .stApp [data-baseweb="input"] > div,
    .stApp [data-baseweb="base-input"] > div,
    .stApp [data-baseweb="textarea"] > div {
        background-color: #242220 !important;
        color: #ffffff !important;
        border-color: #3d3934 !important;
    }

    /* ========================================================
       SELECTBOX
       ======================================================== */

    .stApp [data-baseweb="select"] > div {
        background-color: #242220 !important;
        color: #ffffff !important;
        border-color: #3d3934 !important;
        border-radius: 10px !important;
    }

    .stApp [data-baseweb="select"] span,
    .stApp [data-baseweb="select"] div {
        color: #ffffff !important;
    }

    /* MENIUL SELECTBOX */

    [data-baseweb="popover"],
    [data-baseweb="menu"],
    [role="listbox"] {
        background-color: #242220 !important;
        color: #ffffff !important;
    }

    [role="option"] {
        background-color: #242220 !important;
        color: #ffffff !important;
    }

    [role="option"]:hover,
    [role="option"][aria-selected="true"] {
        background-color: #35312d !important;
        color: #ffffff !important;
    }

    /* ========================================================
       RADIO - CONECTARE / CONT NOU
       ======================================================== */

    .stApp [data-testid="stRadio"] label,
    .stApp [data-testid="stRadio"] p,
    .stApp [data-testid="stRadio"] span {
        color: #ffffff !important;
    }

    /* ========================================================
       BUTOANE
       ======================================================== */

    .stApp .stButton > button,
    .stApp .stFormSubmitButton > button,
    .stApp [data-testid="stFormSubmitButton"] > button {
        background-color: #242220 !important;
        color: #ffffff !important;
        border: 1px solid #3d3934 !important;
        border-radius: 10px !important;
        box-shadow: none !important;
    }

    .stApp .stButton > button:hover,
    .stApp .stFormSubmitButton > button:hover {
        background-color: #2b2926 !important;
        color: #ffffff !important;
        border-color: #5a5149 !important;
    }

    /* BUTON PRINCIPAL LOGIN */

    .stApp [data-testid="stFormSubmitButton"] > button {
        background-color: #d9774a !important;
        color: #ffffff !important;
        font-weight: 600 !important;
        border-color: #d9774a !important;
    }

    .stApp [data-testid="stFormSubmitButton"] > button:hover {
        background-color: #c9683d !important;
        color: #ffffff !important;
    }

    /* ========================================================
       CHAT INPUT
       ======================================================== */

    [data-testid="stBottom"],
    [data-testid="stBottom"] > div,
    [data-testid="stChatInputContainer"],
    .stChatFloatingInputContainer {
        background-color: #1a1917 !important;
        background: #1a1917 !important;
        border: none !important;
    }

    /* ========================================================
       FILE UPLOADER
       ======================================================== */

    .stApp [data-testid="stFileUploaderDropzone"] {
        background-color: #242220 !important;
        border: 1px dashed #3d3934 !important;
        color: #ffffff !important;
    }

    .stApp [data-testid="stFileUploaderDropzone"] * {
        color: #ffffff !important;
    }

    /* ========================================================
       ALERTURI
       ======================================================== */

    .stApp [data-testid="stAlert"] {
        color: #ffffff !important;
    }

    .stApp [data-testid="stAlert"] p,
    .stApp [data-testid="stAlert"] span {
        color: #ffffff !important;
    }

    /* ========================================================
       TABURI
       ======================================================== */

    .stApp button[data-baseweb="tab"] {
        color: #aaa39a !important;
    }

    .stApp button[data-baseweb="tab"][aria-selected="true"] {
        color: #ffffff !important;
    }

    /* ========================================================
       CHAT - MESAJE UTILIZATOR
       ======================================================== */

    .msg-row {
        display: flex;
        margin: 8px 0;
        width: 100%;
    }

    .msg-row.user {
        justify-content: flex-end;
    }

    .msg-row.assistant {
        justify-content: flex-start;
    }

    .bubble-user {
        background-color: #2b2926 !important;
        color: #ffffff !important;
        padding: 12px 18px;
        border-radius: 16px 16px 4px 16px;
        max-width: 82%;
        font-size: 14.5px;
        line-height: 1.55;
        border: 1px solid #3c3935;
    }

    /* ========================================================
       CHAT - MESAJE SCRIBO
       ======================================================== */

    .bubble-assistant-box {
        background-color: #23221f !important;
        border: 1px solid #33302b !important;
        border-radius: 16px 16px 16px 4px;
        padding: 14px 18px;
        margin-bottom: 12px;
        width: 100%;
    }

    .assistant-label {
        font-size: 11.5px;
        font-weight: 700;
        color: #d9774a !important;
        letter-spacing: 0.5px;
        text-transform: uppercase;
        margin-bottom: 6px;
    }

    .bubble-assistant-text {
        color: #f1f5f9 !important;
        font-size: 14.5px;
        line-height: 1.6;
    }

    /* ========================================================
       SIDEBAR
       RĂMÂNE DESCHISĂ, CU FUNDAL DESCHIS ȘI TEXT ÎNCHIS
       ======================================================== */

    section[data-testid="stSidebar"] {
        background-color: #f7f6f2 !important;
        border-right: 1px solid #e5e2da !important;
    }

    section[data-testid="stSidebar"] * {
        color: #2b2823 !important;
    }

    section[data-testid="stSidebar"] .stButton > button {
        text-align: left !important;
        background: transparent !important;
        color: #2b2823 !important;
        border: none !important;
        border-radius: 8px !important;
    }

    section[data-testid="stSidebar"] .stButton > button:hover {
        background: #e9e6df !important;
    }

    section[data-testid="stSidebar"] .stSelectbox > div > div,
    section[data-testid="stSidebar"] [data-baseweb="select"] > div {
        background-color: #ffffff !important;
        color: #2b2823 !important;
        border: 1px solid #ddd8ce !important;
        border-radius: 8px !important;
    }

    section[data-testid="stSidebar"] [data-baseweb="select"] span,
    section[data-testid="stSidebar"] [data-baseweb="select"] div {
        color: #2b2823 !important;
    }

    /* ========================================================
       TITLU SCRIBO
       ======================================================== */

    .hero-title {
        font-size: 28px;
        font-weight: 700;
        text-align: center;
        color: #f1ede7 !important;
        margin-bottom: 12px;
    }

    .hero-spark {
        color: #d9774a !important;
        margin-right: 8px;
    }

    /* ========================================================
       UTILIZATOR
       ======================================================== */

    .user-card {
        background: #ebe8e1;
        color: #2b2823 !important;
        border-radius: 10px;
        padding: 10px 12px;
        margin-top: 8px;
        margin-bottom: 8px;
    }

    .user-card b {
        color: #2b2823 !important;
    }

    .user-avatar {
        display: inline-flex;
        width: 32px;
        height: 32px;
        border-radius: 50%;
        align-items: center;
        justify-content: center;
        background: #d9774a;
        color: white !important;
        font-weight: 700;
        margin-right: 8px;
    }

    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# OIDC LOGIN
# ============================================================

st.markdown(
    """
    <style>

    @import url(
        'https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap'
    );

    html,
    body,
    [class*="css"],
    .stApp {
        background-color: #1a1917 !important;
        color: #e5e5e5 !important;
        font-family: 'Inter', -apple-system, sans-serif;
    }

    footer {
        visibility: hidden;
    }

    header[data-testid="stHeader"] {
        visibility: visible;
        background: transparent !important;
    }

    [data-testid="stBottom"],
    [data-testid="stBottom"] > div,
    [data-testid="stChatInputContainer"],
    .stChatFloatingInputContainer {
        background-color: #1a1917 !important;
        background: #1a1917 !important;
        border: none !important;
    }

    section[data-testid="stSidebar"] {
        background-color: #f7f6f2 !important;
        border-right: 1px solid #e5e2da !important;
    }

    section[data-testid="stSidebar"] * {
        color: #2b2823 !important;
    }

    section[data-testid="stSidebar"] .stButton button {
        text-align: left !important;
    }

    section[data-testid="stSidebar"] .stButton button {
        background: transparent !important;
        border: none !important;
        border-radius: 8px !important;
    }

    section[data-testid="stSidebar"] .stButton button:hover {
        background: #e9e6df !important;
    }

    section[data-testid="stSidebar"] .stSelectbox > div > div {
        background-color: #ffffff !important;
        border: 1px solid #ddd8ce !important;
        border-radius: 8px !important;
    }

    .hero-title {
        font-size: 28px;
        font-weight: 700;
        text-align: center;
        color: #f1ede7;
        margin-bottom: 12px;
    }

    .hero-spark {
        color: #d9774a;
        margin-right: 8px;
    }

    .user-card {
        background: #ebe8e1;
        border-radius: 10px;
        padding: 10px 12px;
        margin-top: 8px;
        margin-bottom: 8px;
    }

    .user-avatar {
        display: inline-flex;
        width: 32px;
        height: 32px;
        border-radius: 50%;
        align-items: center;
        justify-content: center;
        background: #d9774a;
        color: white !important;
        font-weight: 700;
        margin-right: 8px;
    }

    .msg-row {
        display: flex;
        margin: 8px 0;
        width: 100%;
    }

    .msg-row.user {
        justify-content: flex-end;
    }

    .msg-row.assistant {
        justify-content: flex-start;
    }

    .bubble-user {
        background-color: #2b2926 !important;
        color: #ffffff !important;
        padding: 12px 18px;
        border-radius: 16px 16px 4px 16px;
        max-width: 82%;
        font-size: 14.5px;
        line-height: 1.55;
        border: 1px solid #3c3935;
    }

    .bubble-assistant-box {
        background-color: #23221f !important;
        border: 1px solid #33302b !important;
        border-radius: 16px 16px 16px 4px;
        padding: 14px 18px;
        margin-bottom: 12px;
        width: 100%;
    }

    .assistant-label {
        font-size: 11.5px;
        font-weight: 700;
        color: #d9774a !important;
        letter-spacing: 0.5px;
        text-transform: uppercase;
        margin-bottom: 6px;
    }

    .bubble-assistant-text {
        color: #f1f5f9 !important;
        font-size: 14.5px;
        line-height: 1.6;
    }

    .st-key-composer {
        width: min(850px, calc(100vw - 32px));
        margin: 8px auto;
    }

    .st-key-composer .stPopover button {
        width: 36px !important;
        height: 36px !important;
        min-height: 36px !important;
        padding: 0 !important;
        border: 0 !important;
        border-radius: 50% !important;
        background: #34312d !important;
        color: #f5f1ea !important;
        font-size: 18px !important;
    }

    .st-key-composer .stPopover button:hover {
        background: #48423b !important;
    }

    .st-key-composer [data-testid="stHorizontalBlock"] {
        align-items: center;
    }

    .st-key-composer [data-testid="stHorizontalBlock"] button {
        min-height: 32px !important;
        padding: 0 8px !important;
        white-space: nowrap !important;
    }

    div[data-testid="stChatInput"] {
        max-width: 850px !important;
        margin: 0 auto !important;
    }

    div[data-testid="stChatInput"] > div,
    div[data-testid="stChatInput"] [data-baseweb="base-input"],
    div[data-testid="stChatInput"] [data-baseweb="input"],
    div[data-testid="stChatInput"] [data-baseweb="textarea"] {
        min-height: 92px !important;
        background-color: #201f1c !important;
        border: 1px solid #413c35 !important;
        border-radius: 24px !important;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.22) !important;
    }

    div[data-testid="stChatInput"] textarea {
        min-height: 88px !important;
        padding: 17px 58px 46px 104px !important;
        color: #ffffff !important;
        -webkit-text-fill-color: #ffffff !important;
        caret-color: #d9774a !important;
    }

    div[data-testid="stChatInput"] button {
        background-color: #d9774a !important;
        color: #ffffff !important;
    }

    /* ========================================================
       ACCESIBILITATE + RESPONSIVE
       ======================================================== */

    .main .block-container {
        max-width: 1180px;
        padding-left: clamp(12px, 3vw, 38px) !important;
        padding-right: clamp(12px, 3vw, 38px) !important;
        padding-bottom: 130px !important;
    }

    .login-shell {
        max-width: 520px;
        margin: 0 auto;
    }

    button, input, textarea, select,
    [role="button"], [role="tab"], [role="radio"] {
        min-height: 44px;
    }

    button:focus-visible,
    input:focus-visible,
    textarea:focus-visible,
    [role="button"]:focus-visible,
    [role="tab"]:focus-visible,
    [role="radio"]:focus-visible {
        outline: 3px solid #f4b183 !important;
        outline-offset: 3px !important;
    }

    [data-testid="stSidebar"] button,
    [data-testid="collapsedControl"] button,
    [data-testid="stSidebarCollapsedControl"] button {
        min-width: 44px !important;
        min-height: 44px !important;
    }

    .bubble-user,
    .bubble-assistant-box {
        overflow-wrap: anywhere;
        word-break: normal;
    }

    @media (prefers-reduced-motion: reduce) {
        *, *::before, *::after {
            scroll-behavior: auto !important;
            transition: none !important;
            animation: none !important;
        }
    }

    @media (max-width: 900px) {
        .bubble-user { max-width: 92%; }
        .bubble-assistant-box { width: 100%; }
    }

    @media (max-width: 640px) {
        .st-key-composer {
            width: calc(100vw - 20px);
            margin: 6px auto;
        }

        .st-key-composer [data-testid="stHorizontalBlock"] {
            gap: 8px;
            flex-wrap: wrap !important;
        }

        .st-key-composer [data-testid="stHorizontalBlock"] button {
            font-size: 14px !important;
            padding: 0 10px !important;
            min-height: 44px !important;
        }

        div[data-testid="stChatInput"] textarea {
            min-height: 76px !important;
            padding: 14px 52px 38px 14px !important;
            font-size: 16px !important;
        }

        div[data-testid="stChatInput"] > div,
        div[data-testid="stChatInput"] [data-baseweb="base-input"],
        div[data-testid="stChatInput"] [data-baseweb="input"],
        div[data-testid="stChatInput"] [data-baseweb="textarea"] {
            min-height: 80px !important;
            border-radius: 18px !important;
        }

        .st-key-composer [data-testid="stHorizontalBlock"]:has(
            [data-testid="column"]:nth-child(5)
        ) {
            flex-wrap: wrap;
            overflow-x: visible;
        }
    }

    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# OIDC LOGIN
# ============================================================

handle_oidc_login()
enhance_accessibility_and_follow_chat(scroll_to_bottom=False)


# ============================================================
# LOGIN / ÎNREGISTRARE
# ============================================================

if not st.session_state.authenticated:

    st.markdown('<div class="login-shell">', unsafe_allow_html=True)
    st.image(
        "assets/logo_scribo.png",
        width=240,
    )

    with st.container():

        auth_mode = st.radio(
            "Alege operațiunea",
            [
                "Conectare",
                "Cont Nou",
            ],
            horizontal=True,
            label_visibility="collapsed",
        )

        # ----------------------------------------------------
        # LOGIN EXTERN
        # ----------------------------------------------------

        st.markdown("### Conectare rapidă")
        st.caption("Poți folosi un cont extern sau contul local Scribo.")

        google_ready = oidc_provider_configured("google")
        apple_ready = oidc_provider_configured("apple")

        if google_ready:
            st.button(
                "Continuă cu Google",
                key="login_google",
                use_container_width=True,
                on_click=start_oidc_login,
                args=("google",),
                help="Autentificare securizată prin Google OpenID Connect",
            )
        else:
            st.button(
                "Continuă cu Google",
                key="login_google_disabled",
                disabled=True,
                use_container_width=True,
                help="Lipsește configurația completă [auth.google] din secrets.toml",
            )

        if apple_ready:
            st.button(
                "Continuă cu Apple",
                key="login_apple",
                use_container_width=True,
                on_click=start_oidc_login,
                args=("apple",),
                help="Autentificare securizată prin Sign in with Apple / OIDC",
            )
        else:
            st.button(
                "Continuă cu Apple",
                key="login_apple_disabled",
                disabled=True,
                use_container_width=True,
                help="Lipsește configurația completă [auth.apple] din secrets.toml",
            )

        st.divider()

        # ----------------------------------------------------
        # CONT LOCAL
        # ----------------------------------------------------

        with st.form("auth_form"):

            if auth_mode == "Cont Nou":

                username_input = st.text_input(
                    "Nume utilizator",
                    placeholder="Exemplu: gom",
                )

                email_input = st.text_input(
                    "Email",
                    placeholder="exemplu@email.com",
                )

                password_input = st.text_input(
                    "Parolă",
                    type="password",
                    placeholder="Minimum 6 caractere",
                )

                password_confirm = st.text_input(
                    "Confirmă parola",
                    type="password",
                )

                submitted = st.form_submit_button(
                    "Creează cont",
                    use_container_width=True,
                )

                if submitted:

                    if password_input != password_confirm:
                        st.error(
                            "Parolele nu coincid."
                        )

                    else:
                        ok, message = register_user(
                            username_input,
                            email_input,
                            password_input,
                        )

                        if ok:
                            st.success(
                                "Cont creat! "
                                "Acum te poți conecta."
                            )
                        else:
                            st.error(message)

            else:

                username_input = st.text_input(
                    "Email sau nume utilizator",
                )

                password_input = st.text_input(
                    "Parolă",
                    type="password",
                )

                submitted = st.form_submit_button(
                    "Intră în cont",
                    use_container_width=True,
                )

                if submitted:

                    user_data = login_user(
                        username_input,
                        password_input,
                    )

                    if user_data:

                        st.session_state.authenticated = True
                        st.session_state.user = user_data[
                            "username"
                        ]
                        st.session_state.username = user_data[
                            "username"
                        ]
                        st.session_state.email = user_data[
                            "email"
                        ]
                        st.session_state.display_name = user_data[
                            "display_name"
                        ]
                        st.session_state.style = user_data[
                            "style"
                        ]
                        st.session_state.voice_model = user_data[
                            "voice"
                        ]

                        # IMPORTANT:
                        # prima intrare = conversație nouă
                        if not st.session_state.auth_initialized:
                            start_new_conversation()
                            st.session_state.auth_initialized = True

                        st.rerun()

                    else:
                        st.error(
                            "Email/nume utilizator sau "
                            "parolă incorectă."
                        )

    st.markdown("</div>", unsafe_allow_html=True)

    st.stop()


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:

    # --------------------------------------------------------
    # BRANDING
    # --------------------------------------------------------

    sidebar_logo_cols = st.columns([1, 2, 1])

    with sidebar_logo_cols[1]:

        st.image(
            "assets/logo_scribo.png",
            width=180,
        )

    st.caption(
        "Studio inteligent pentru scriere."
    )

    # --------------------------------------------------------
    # CONVERSAȚIE NOUĂ
    # --------------------------------------------------------

    if st.button(
        "✏️  Conversație nouă",
        use_container_width=True,
    ):
        start_new_conversation()
        st.rerun()

    st.divider()

    # --------------------------------------------------------
    # CONVERSAȚII
    # --------------------------------------------------------

    st.markdown("#### Conversații")

    conversations = get_conversations(
        st.session_state.username
    )

    if conversations:

        for conversation in conversations[:30]:

            active = (
                conversation["id"]
                == st.session_state.conversation_id
            )

            label = conversation["title"]

            if active:
                label = "● " + label

            if st.button(
                label,
                key=f"conversation_{conversation['id']}",
                use_container_width=True,
            ):
                load_conversation(
                    conversation["id"]
                )
                st.rerun()

    else:
        st.caption(
            "Nu există conversații salvate."
        )

    st.divider()

    # --------------------------------------------------------
    # PROIECTE
    # --------------------------------------------------------

    st.markdown("#### 📁 Proiecte")

    projects = get_projects(
        st.session_state.username
    )

    if projects:

        project_names = [
            p["name"]
            for p in projects
        ]

        current_project_name = None

        if st.session_state.project_id:
            current_project = get_project(
                st.session_state.project_id,
                st.session_state.username,
            )

            if current_project:
                current_project_name = current_project[
                    "name"
                ]

        if current_project_name not in project_names:
            current_project_name = project_names[0]

        selected_project_name = st.selectbox(
            "Proiect",
            project_names,
            index=project_names.index(
                current_project_name
            ),
            label_visibility="collapsed",
        )

        selected_project = next(
            p
            for p in projects
            if p["name"] == selected_project_name
        )

        if selected_project["id"] != st.session_state.project_id:

            st.session_state.project_id = selected_project[
                "id"
            ]

            st.session_state.roman = selected_project[
                "roman"
            ] or ""

            st.session_state.style = selected_project[
                "style"
            ] or "Fantezie Alchimie"

            st.rerun()

    if st.button(
        "＋ Proiect nou",
        use_container_width=True,
    ):

        new_project_id = create_project(
            st.session_state.username,
            f"Proiect {len(projects) + 1}",
            st.session_state.style,
        )

        st.session_state.project_id = new_project_id
        st.session_state.roman = ""
        st.session_state.messages = []

        new_conversation_id = create_conversation(
            st.session_state.username,
            new_project_id,
            "Conversație nouă",
        )

        st.session_state.conversation_id = (
            new_conversation_id
        )

        st.rerun()

    st.divider()

    # --------------------------------------------------------
    # STIL
    # --------------------------------------------------------

    st.markdown("#### 📚 Stil literar")

    style_options = [
        "Fantezie Alchimie",
        "Epic",
        "Dark Fantasy",
        "Misterios",
        "Romantic",
        "Realist",
        "SF",
        "Thriller",
    ]

    current_style_idx = (
        style_options.index(
            st.session_state.style
        )
        if st.session_state.style
        in style_options
        else 0
    )

    new_style = st.selectbox(
        "Stil literar",
        style_options,
        index=current_style_idx,
        label_visibility="collapsed",
    )

    if new_style != st.session_state.style:

        st.session_state.style = new_style

        if st.session_state.project_id:
            update_project(
                st.session_state.project_id,
                st.session_state.username,
                st.session_state.roman,
                new_style,
            )

    st.divider()

    # --------------------------------------------------------
    # VOCE
    # --------------------------------------------------------

    st.markdown("#### 🔊 Voce AI")

    voice_map = {
        "Alina (Feminin)": "ro-RO-AlinaNeural",
        "Emil (Masculin)": "ro-RO-EmilNeural",
    }

    current_voice_name = [
        key
        for key, value in voice_map.items()
        if value == st.session_state.voice_model
    ]

    selected_voice_name = (
        current_voice_name[0]
        if current_voice_name
        else "Alina (Feminin)"
    )

    new_voice_label = st.selectbox(
        "Selectează vocea",
        list(voice_map.keys()),
        index=list(
            voice_map.keys()
        ).index(selected_voice_name),
        label_visibility="collapsed",
    )

    new_voice_val = voice_map[
        new_voice_label
    ]

    if new_voice_val != st.session_state.voice_model:

        st.session_state.voice_model = new_voice_val

        conn = get_conn()

        conn.execute(
            """
            UPDATE users
            SET voice = ?
            WHERE username = ?
            """,
            (
                new_voice_val,
                st.session_state.username,
            ),
        )

        conn.commit()
        conn.close()

    st.divider()

    # --------------------------------------------------------
    # ROMAN
    # --------------------------------------------------------

    st.markdown("#### 📖 Romanul tău")

    words = (
        len(
            st.session_state.roman.split()
        )
        if st.session_state.roman.strip()
        else 0
    )

    st.caption(
        f"{words:,} / 50.000 cuvinte"
        .replace(",", ".")
    )

    st.progress(
        min(words / 50000, 1.0)
    )

    st.divider()

    # --------------------------------------------------------
    # UTILIZATOR
    # --------------------------------------------------------

    username = (
        st.session_state.username
        or "utilizator"
    )

    display_name = (
        st.session_state.display_name
        or username
    )

    initial = (
        username[:1].upper()
        if username
        else "U"
    )

    st.markdown(
        f"""
        <div class="user-card">
            <span class="user-avatar">
                {escape_html(initial)}
            </span>
            <b>{escape_html(username)}</b>
            <br>
            <span style="
                font-size:11px;
                margin-left:44px;
                color:#777 !important;
            ">
                {escape_html(display_name)}
            </span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if st.button(
        "🚪 Deconectare",
        use_container_width=True,
    ):

        # Logout OIDC dacă este activ
        if oidc_available():
            try:
                if st.user.is_logged_in:
                    st.logout()
            except Exception:
                pass

        st.session_state.authenticated = False
        st.session_state.user = None
        st.session_state.username = None
        st.session_state.email = ""
        st.session_state.display_name = ""
        st.session_state.messages = []
        st.session_state.roman = ""
        st.session_state.conversation_id = None
        st.session_state.project_id = None
        st.session_state.auth_initialized = False

        st.rerun()


# ============================================================
# ZONA PRINCIPALĂ
# ============================================================

logo_cols = st.columns([1, 1, 1])

with logo_cols[1]:

    st.image(
        "assets/logo_scribo.png",
        width=280,
    )

_, center_col, _ = st.columns(
    [0.1, 0.8, 0.1]
)


with center_col:

    chat_tab, roman_tab = st.tabs(
        ["💬 Chat", "📖 Roman"]
    )

    # ========================================================
    # CHAT
    # ========================================================

    with chat_tab:

        if not st.session_state.messages:

            st.markdown(
                """
                <div style="
                    text-align:center;
                    color:#88847c;
                    margin-top:60px;
                    margin-bottom:60px;
                ">
                    <h4>
                        Spune-mi ce dorești să scriem astăzi! 🚀
                    </h4>
                    <p style="font-size:13.5px;">
                        Tastează un mesaj sau alege
                        o opțiune de mai jos.
                    </p>
                </div>
                """,
                unsafe_allow_html=True,
            )

        # ----------------------------------------------------
        # MESAJE
        # ----------------------------------------------------

        for i, message in enumerate(
            st.session_state.messages
        ):

            role = message["role"]
            content = message["content"]

            if role == "user":

                images = message.get(
                    "images",
                    [],
                )

                st.markdown(
                    '<div class="msg-row user">',
                    unsafe_allow_html=True,
                )

                bubble_html = (
                    '<div class="bubble-user">'
                )

                if images:

                    bubble_html += (
                        '<div style="'
                        'display:flex;'
                        'gap:6px;'
                        'margin-bottom:6px;">'
                    )

                    for img in images:

                        try:
                            uri = image_to_data_uri(
                                img["data"]
                            )

                            bubble_html += (
                                f'<img src="{uri}" '
                                'style="'
                                'width:70px;'
                                'height:70px;'
                                'border-radius:6px;'
                                'object-fit:cover;" />'
                            )
                        except Exception:
                            pass

                    bubble_html += "</div>"

                bubble_html += (
                    escape_html(content)
                    .replace("\n", "<br>")
                    + "</div></div>"
                )

                st.markdown(
                    bubble_html,
                    unsafe_allow_html=True,
                )

            else:

                st.markdown(
                    '<div class="bubble-assistant-box">',
                    unsafe_allow_html=True,
                )

                h_left, h_right = st.columns(
                    [10, 1]
                )

                with h_left:

                    st.markdown(
                        """
                        <div class="assistant-label">
                            🖋️ SCRIBO
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                with h_right:

                    with st.popover(
                        "⋮",
                        use_container_width=False,
                    ):

                        st.download_button(
                            "⬇️ TXT",
                            data=content,
                            file_name=f"scribo_{i}.txt",
                            key=f"txt_{i}",
                            use_container_width=True,
                        )

                        if st.button(
                            "📖 În Roman",
                            key=f"add_{i}",
                            use_container_width=True,
                        ):

                            st.session_state.roman = (
                                (
                                    st.session_state.roman
                                    + "\n\n"
                                    + content.strip()
                                )
                                .strip()
                            )

                            if st.session_state.project_id:

                                update_project(
                                    st.session_state.project_id,
                                    st.session_state.username,
                                    st.session_state.roman,
                                    st.session_state.style,
                                )

                            st.success(
                                "Adăugat în roman!"
                            )

                safe_content = escape_html(
                    clean_ai_text(content)
                ).replace(
                    "\n",
                    "<br>",
                )

                st.markdown(
                    f"""
                    <div class="bubble-assistant-text">
                        {safe_content}
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                st.markdown(
                    "</div>",
                    unsafe_allow_html=True,
                )

                if (
                    "audio" in message
                    and message["audio"]
                ):
                    st.audio(
                        message["audio"],
                        format="audio/mp3",
                    )

        # ----------------------------------------------------
        # CONTINUARE TEXT
        # ----------------------------------------------------

        if (
            st.session_state.messages
            and st.session_state.messages[-1]["role"]
            == "assistant"
        ):

            last_text = (
                st.session_state.messages[-1]["content"]
                .strip()
            )

            action_cols = st.columns(4)

            with action_cols[0]:
                if st.button("Vreau mai mult", use_container_width=True):
                    st.session_state.pending_prompt = (
                        "Dezvoltă răspunsul anterior cu mai multe detalii relevante."
                    )
                    st.rerun()

            with action_cols[1]:
                if st.button("Extinde", use_container_width=True):
                    st.session_state.pending_prompt = (
                        "Extinde răspunsul anterior, păstrând ideile și stilul."
                    )
                    st.rerun()

            with action_cols[2]:
                if st.button("Reîncearcă", use_container_width=True):
                    st.session_state.pending_prompt = (
                        "Reformulează răspunsul anterior într-un mod mai clar și mai convingător."
                    )
                    st.rerun()

            with action_cols[3]:
                if st.button("Treci la următoarea analiză", use_container_width=True):
                    st.session_state.pending_prompt = (
                        "Treci la următoarea analiză relevantă pentru cererea utilizatorului."
                    )
                    st.rerun()

        st.markdown(
            '<div id="chat-bottom-anchor" '
            'style="height:1px;"></div>',
            unsafe_allow_html=True,
        )
        enhance_accessibility_and_follow_chat(
            scroll_to_bottom=bool(st.session_state.messages)
        )

        # ----------------------------------------------------
        # ATAȘAMENTE
        # ----------------------------------------------------

        st.markdown(
            '<div class="composer-wrap">',
            unsafe_allow_html=True,
        )

        if (
            st.session_state.attached_file_name
            or st.session_state.attached_images
        ):

            chip_cols = st.columns(
                [6, 1]
            )

            with chip_cols[0]:

                chips = ""

                if st.session_state.attached_file_name:

                    chips += (
                        '<span style="'
                        'background:#312e29;'
                        'padding:4px 8px;'
                        'border-radius:6px;'
                        'font-size:12px;'
                        'margin-right:6px;">'
                        f'📄 {escape_html(st.session_state.attached_file_name)}'
                        '</span>'
                    )

                for img in st.session_state.attached_images:

                    chips += (
                        '<span style="'
                        'background:#312e29;'
                        'padding:4px 8px;'
                        'border-radius:6px;'
                        'font-size:12px;'
                        'margin-right:6px;">'
                        f'🖼️ {escape_html(img["name"])}'
                        '</span>'
                    )

                st.markdown(
                    chips,
                    unsafe_allow_html=True,
                )

            with chip_cols[1]:

                if st.button(
                    "Elimină atașamentele",
                    key="clear_attachments",
                    use_container_width=True,
                ):

                    st.session_state.attached_text = ""
                    st.session_state.attached_file_name = ""
                    st.session_state.attached_images = []

                    st.rerun()

# ============================================================
# BUTOANE + ATAȘAMENTE + VOCE
# ============================================================

render_suggestions()

composer = st.container(key="composer")

with composer:

    prompt = st.chat_input(
        "Scrie mesajul pentru Scribo",
        key="main_chat_input",
    )

    if st.session_state.attached_text and not prompt:
        if st.button(
            "Trimite documentul atașat",
            key="send_attached_document",
            use_container_width=True,
        ):
            prompt = "Analizează documentul atașat."

    composer_actions = st.columns(2)

    # ========================================================
    # ATAȘAMENTE
    # ========================================================

    with composer_actions[0]:

        with st.expander(
            "Atașează fișier sau imagine",
            expanded=False,
        ):

            doc_tab, img_tab = st.tabs(
                [
                    "📄 Document",
                    "🖼️ Imagine",
                ]
            )

            # ------------------------------------------------
            # DOCUMENT
            # ------------------------------------------------

            with doc_tab:

                doc_file = st.file_uploader(
                    "Document",
                    type=[
                        "txt",
                        "pdf",
                        "docx",
                    ],
                    key="doc_uploader",
                    label_visibility="collapsed",
                )

                if doc_file is not None:

                    extracted = extract_document_text(
                        doc_file
                    )

                    if extracted:

                        st.session_state.attached_text = extracted

                        st.session_state.attached_file_name = (
                            doc_file.name
                        )

                        st.success(
                            "Document atașat!"
                        )

            # ------------------------------------------------
            # IMAGINE
            # ------------------------------------------------

            with img_tab:

                img_files = st.file_uploader(
                    "Imagini",
                    type=[
                        "png",
                        "jpg",
                        "jpeg",
                        "webp",
                    ],
                    accept_multiple_files=True,
                    key="img_uploader",
                    label_visibility="collapsed",
                )

                if img_files:

                    st.session_state.attached_images = [
                        {
                            "name": f.name,
                            "data": f.getvalue(),
                        }
                        for f in img_files
                    ]

                    st.success(
                        "Imagini atașate!"
                    )

    # ========================================================
    # MICROFON
    # ========================================================

    with composer_actions[1]:

        with st.expander(
            "Mesaj vocal",
            expanded=False,
        ):

            st.markdown(
                "**Înregistrează vocea**"
            )

            audio = st.audio_input(
                "Vorbește acum",
                key="voice_recorder",
                label_visibility="collapsed",
            )

            if audio:

                try:

                    with st.spinner(
                        "Transcriu..."
                    ):

                        transcription = (
                            client.audio.transcriptions.create(
                                file=(
                                    "audio.wav",
                                    audio.getvalue(),
                                ),
                                model=WHISPER_MODEL,
                                language="ro",
                            )
                        )

                    voice_text = getattr(
                        transcription,
                        "text",
                        str(transcription),
                    ).strip()

                    if voice_text:

                        st.info(
                            f"🎙️ {voice_text}"
                        )

                        if st.button(
                            "Trimite mesaj vocal",
                            key="send_voice_now",
                        ):

                            st.session_state.pending_prompt = (
                                voice_text
                            )

                            st.session_state.is_voice_input = (
                                True
                            )

                            st.rerun()

                except Exception as e:

                    st.error(
                        f"Eroare transcriere: {e}"
                    )
        st.markdown(
            "</div>",
            unsafe_allow_html=True,
        )

    # ========================================================
    # ROMAN
    # ========================================================

    with roman_tab:

        st.subheader(
            "📖 Romanul tău"
        )

        roman_text = st.text_area(
            "Editor text",
            value=st.session_state.roman,
            height=500,
            placeholder=(
                "Scrie sau editează "
                "povestea ta aici..."
            ),
            label_visibility="collapsed",
        )

        if roman_text != st.session_state.roman:

            st.session_state.roman = roman_text

            if st.session_state.project_id:

                update_project(
                    st.session_state.project_id,
                    st.session_state.username,
                    st.session_state.roman,
                    st.session_state.style,
                )

        col1, col2 = st.columns(2)

        with col1:

            st.download_button(
                "⬇️ Descarcă TXT",
                data=st.session_state.roman,
                file_name="roman.txt",
                mime="text/plain",
                use_container_width=True,
            )

        with col2:

            if DOCX_AVAILABLE:

                doc = Document()

                doc.add_paragraph(
                    st.session_state.roman
                )

                buf = BytesIO()

                doc.save(buf)

                buf.seek(0)

                st.download_button(
                    "📄 Descarcă DOCX",
                    data=buf,
                    file_name="roman.docx",
                    mime=(
                        "application/vnd.openxmlformats-"
                        "officedocument.wordprocessingml.document"
                    ),
                    use_container_width=True,
                )


# ============================================================
# TRANSMITERE MESAJ GROQ
# ============================================================

if st.session_state.pending_prompt:

    prompt = (
        st.session_state.pending_prompt
    )

    st.session_state.pending_prompt = None


if prompt:

    prompt = prompt.strip()

    if prompt:

        is_voice_turn = (
            st.session_state.is_voice_input
        )

        st.session_state.is_voice_input = False

        # ----------------------------------------------------
        # CĂUTARE WEB
        # ----------------------------------------------------

        web_context = ""

        if needs_web_search(prompt):

            with st.spinner(
                "🔍 Caut informații verificate pe internet..."
            ):

                found_info = get_web_context(
                    prompt
                )

                if found_info:

                    web_context = (
                        "\n\n"
                        "[REZULTATE CĂUTARE WEB "
                        "PENTRU VERIFICARE]:\n"
                        f"{found_info}\n"
                    )

        # ----------------------------------------------------
        # CONVERSAȚIE
        # ----------------------------------------------------

        if not st.session_state.conversation_id:

            start_new_conversation()

        conversation_id = (
            st.session_state.conversation_id
        )

        current_images = list(
            st.session_state.attached_images
        )

        # ----------------------------------------------------
        # PRIMUL MESAJ = TITLU
        # ----------------------------------------------------

        if not st.session_state.messages:

            title = make_conversation_title(
                prompt
            )

            update_conversation_title(
                conversation_id,
                st.session_state.username,
                title,
            )

        # ----------------------------------------------------
        # SALVĂM MESAJUL USERULUI
        # ----------------------------------------------------

        add_message(
            conversation_id,
            st.session_state.username,
            "user",
            prompt,
            current_images,
        )

        st.session_state.messages.append(
            {
                "role": "user",
                "content": prompt,
                "images": current_images,
            }
        )

        st.session_state.attached_images = []

        # ----------------------------------------------------
        # CONTEXT ROMAN
        # ----------------------------------------------------

        roman_context = ""

        if st.session_state.roman.strip():

            roman_context = (
                "\n\n"
                "[FRAGMENT DIN ROMANUL "
                "UTILIZATORULUI]:\n"
                f"{st.session_state.roman[-2000:]}\n"
            )

        # ----------------------------------------------------
        # DOCUMENT
        # ----------------------------------------------------

        document_context = ""

        if st.session_state.attached_text:

            document_context = (
                "\n\n"
                "[DOCUMENT ATAȘAT]:\n"
                f"{st.session_state.attached_text}\n"
            )

        # ----------------------------------------------------
        # SISTEM
        # ----------------------------------------------------

        system_prompt = f"""
Ești SCRIBO AI, un asistent literar de elită
specializat în scriere creativă, editare și
consultanță narativă.

Utilizatorul curent:
{st.session_state.username}

Stil literar curent:
{st.session_state.style}

REGULI STRICTE DE SIGURANȚĂ ȘI PRECIZIE FACTUALĂ:

1. VERIFICĂRI FACTUALE

Dacă utilizatorul cere detalii despre:
- cărți existente;
- autori;
- date de lansare;
- evenimente istorice;
- premii literare;

folosește exclusiv informațiile factuale
prezente în secțiunea:

[REZULTATE CĂUTARE WEB PENTRU VERIFICARE]

Dacă nu există informații verificate,
spune sincer că nu știi.

Este strict interzis să inventezi:
- autori reali;
- cărți reale;
- edituri reale;
- date istorice;
- personaje reale;
- intrigi ale unor cărți existente.

2. CREATIVITATE

Când utilizatorul cere ficțiune,
scene originale, personaje, dialoguri
sau idei, aplică stilul:

{st.session_state.style}

3. CONTINUITATE

Ține cont de fragmentul romanului
și de documentul atașat atunci când
sunt relevante.

4. FORMAT RĂSPUNS

Scrie întotdeauna text simplu, curat,
ca într-un mesaj tastat de o persoană.
Nu folosi Markdown, asteriscuri, diezuri,
liste cu marcatori, cod sau linkuri formatate.

{roman_context}

{document_context}

{web_context}
"""

        # ----------------------------------------------------
        # MESAJE RECENTE
        # ----------------------------------------------------

        recent_messages = (
            st.session_state.messages[-8:]
        )

        api_messages = [
            {
                "role": "system",
                "content": system_prompt,
            }
        ]

        for msg in recent_messages:

            api_messages.append(
                {
                    "role": msg["role"],
                    "content": msg["content"],
                }
            )

        # ----------------------------------------------------
        # GROQ
        # ----------------------------------------------------

        try:

            with st.spinner(
                "🖋️ Scribo scrie..."
            ):

                response = (
                    client.chat.completions.create(
                        model=MODEL_GROQ,
                        messages=api_messages,
                        temperature=0.6,
                        max_completion_tokens=1500,
                        top_p=0.95,
                        include_reasoning=False,
                    )
                )

            answer = clean_ai_text(
                response.choices[0]
                .message.content
                or ""
            )

            if answer:

                audio_bytes = None

                if (
                    is_voice_turn
                    and EDGE_TTS_AVAILABLE
                ):

                    with st.spinner(
                        "🔊 Sintetizez vocea..."
                    ):

                        audio_bytes = (
                            get_audio_response(
                                answer,
                                st.session_state.voice_model,
                            )
                        )

                # ------------------------------------------------
                # SALVARE DB
                # ------------------------------------------------

                add_message(
                    conversation_id,
                    st.session_state.username,
                    "assistant",
                    answer,
                    [],
                )

                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": answer,
                        "audio": audio_bytes,
                    }
                )

                st.rerun()

        except Exception as e:

            st.error(
                f"Eroare la comunicarea cu serverul Groq: {e}"
            )


# ============================================================
# AUTO-SCROLL
# ============================================================

components.html(
    """
    <script>

    try {
        if (window.top.location.search.includes("utm_medium=oembed")) {
            window.top.location.replace(
                window.top.location.pathname + "?embed=true"
            );
        }
    } catch (e) {}

    function doScroll() {

        try {

            const parentDoc =
                window.parent.document;

            const anchor =
                parentDoc.getElementById(
                    "chat-bottom-anchor"
                );

            if (anchor) {

                anchor.scrollIntoView({
                    behavior: "smooth",
                    block: "end"
                });
            }

            const mainSection =
                parentDoc.querySelector(
                    "section.main"
                );

            if (mainSection) {

                mainSection.scrollTo({
                    top: mainSection.scrollHeight,
                    behavior: "smooth"
                });
            }

            const hidePlatformWidgets = () => {
                const selectors = [
                    '[data-testid="manage-app-button"]',
                    'iframe[src*="statuspage.io"]',
                    '.viewerBadge_container__1QS13',
                    '[data-testid="stStatusWidget"]'
                ];

                for (const getDoc of [
                    () => parentDoc,
                    () => window.top.document
                ]) {
                    try {
                        const doc = getDoc();
                        selectors.forEach((selector) => {
                            doc.querySelectorAll(selector).forEach((element) => {
                                element.style.display = "none";
                            });
                        });
                    } catch (e) {}
                }
            };

            hidePlatformWidgets();
            setInterval(hidePlatformWidgets, 1000);

            window.parent.scrollTo({
                top: parentDoc.body.scrollHeight,
                behavior: "smooth"
            });

        } catch (e) {}

    }

    setTimeout(doScroll, 100);
    setTimeout(doScroll, 400);
    setTimeout(doScroll, 800);

    </script>
    """,
    height=0,
    width=0,
)